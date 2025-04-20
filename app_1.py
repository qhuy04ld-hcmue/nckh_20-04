import os
from flask import Flask, request, render_template
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch
import docx
from google.cloud import vision
import io
from googleapiclient.discovery import build
from vncorenlp import VnCoreNLP

# Đường dẫn tới mô hình VnCoreNLP (tải từ GitHub)
vncorenlp = VnCoreNLP("F:/VnCoreNLP-master/VnCoreNLP-1.1.1.jar")  # Cập nhật đúng đường dẫn

#YOUTUBE_API_KEY = "YOUR_API_KEY"  # <<< thay bằng key thật của bạn
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "./api_gg.json"

app = Flask(__name__)
model = AutoModelForSequenceClassification.from_pretrained("saved_model")
tokenizer = AutoTokenizer.from_pretrained("saved_model")

def classify_text(text):
    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True, max_length=128)
    with torch.no_grad():
        outputs = model(**inputs)
        probs = torch.nn.functional.softmax(outputs.logits, dim=-1)

        print("🔍 Xác suất phân lớp:", probs.tolist())  # In xác suất ra terminal/log

        label = torch.argmax(probs).item()
        confidence = probs[0][label].item()
    return ("Vật lý" if label == 0 else "Hóa học", round(confidence * 100, 2))



def extract_text_from_file(file):
    if file.filename.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def extract_text_from_image(image):
    try:
        client = vision.ImageAnnotatorClient()
        content = image.read()
        image = vision.Image(content=content)
        response = client.text_detection(image=image)
        texts = response.text_annotations
        return texts[0].description if texts else ""
    except Exception as e:
        print("Vision API Error:", e)
        return ""

# Hàm trích xuất từ khóa bằng vncorenlp
def extract_keywords_vncorenlp(text):
    sentences = vncorenlp.tokenize(text)  # Tokenize văn bản
    pos_tags = vncorenlp.pos_tag(text)  # POS tagging
    
    # Lọc danh từ và động từ
    keywords = [word for word, tag in zip(sentences[0], pos_tags[0]) if tag in ['N', 'V'] and len(word) > 2]

    return keywords[:5]  # Chỉ lấy 5 từ khóa đầu tiên

def search_youtube(keywords):
    youtube = build("youtube", "v3", developerKey=YOUTUBE_API_KEY)
    query = " ".join(keywords)
    request = youtube.search().list(
        q=query,
        part="snippet",
        maxResults=25,
        type="video"
    )
    response = request.execute()
    videos = []
    for item in response["items"]:
        video = {
            "title": item["snippet"]["title"],
            "link": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
        }
        videos.append(video)
    return videos

@app.route("/", methods=["GET", "POST"])
def index():
    result, videos = None, []
    if request.method == "POST":
        text = ""
        if "input_text" in request.form and request.form["input_text"]:
            text = request.form["input_text"]
        elif "file" in request.files and request.files["file"]:
            text = extract_text_from_file(request.files["file"])
        elif "image" in request.files and request.files["image"]:
            text = extract_text_from_image(request.files["image"])

        if text:
            label, confidence = classify_text(text)
            keywords = extract_keywords_vncorenlp(text)  # Sử dụng hàm trích xuất từ khóa từ vncorenlp
            videos = search_youtube(keywords)
            result = {
                "text": text,
                "label": label,
                "confidence": confidence,
                "keywords": keywords
            }

    return render_template("index.html", result=result, videos=videos)

if __name__ == "__main__":
    app.run(debug=True)










"""
import os
from flask import Flask, request, render_template
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import torch
import docx
from google.cloud import vision
import io
#from youtubesearchpython import VideosSearch
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from googleapiclient.discovery import build

from stopwords import get_stopwords



YOUTUBE_API_KEY = "YOUR_API_KEY"  # <<< thay bằng key thật của bạn
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "./api_gg.json"



app = Flask(__name__)
model = AutoModelForSequenceClassification.from_pretrained("saved_model")
tokenizer = AutoTokenizer.from_pretrained("saved_model")

def classify_text(text):
    inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True, max_length=128)
    with torch.no_grad():
        outputs = model(**inputs)
        probs = torch.nn.functional.softmax(outputs.logits, dim=-1)
        label = torch.argmax(probs).item()
        confidence = probs[0][label].item()
    return ("Vật lý" if label == 0 else "Hóa học", round(confidence * 100, 2))

def extract_text_from_file(file):
    if file.filename.endswith(".txt"):
        return file.read().decode("utf-8")
    elif file.filename.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def extract_text_from_image(image):
    try:
        client = vision.ImageAnnotatorClient()
        content = image.read()
        image = vision.Image(content=content)
        response = client.text_detection(image=image)
        texts = response.text_annotations
        return texts[0].description if texts else ""
    except Exception as e:
        print("Vision API Error:", e)
        return ""


def extract_keywords1(text):
    words = word_tokenize(text.lower())
    filtered = [w for w in words if w.isalnum() and w not in stopwords.words("english") and len(w) > 2]
    return list(set(filtered))[:5]  # 5 keywords

stopwords_vn = set(get_stopwords("vi"))

def extract_keywords(text):
    # Tokenize văn bản tiếng Việt
    words = word_tokenize(text, format="text").split()

    # Lọc bỏ stopwords tiếng Việt
    filtered = [w for w in words if w.isalnum() and w not in stopwords_vn and len(w) > 2]

    # Lấy ra 5 từ khóa đầu tiên
    return list(set(filtered))[:5]  # Giới hạn 5 từ khóa
'''
def search_youtube(keywords):
    query = " ".join(keywords)
    videos_search = VideosSearch(query, limit=27)
    results = videos_search.result()["result"]
    return [{"title": v["title"], "link": v["link"]} for v in results]
'''
def search_youtube(keywords):
    youtube = build("youtube", "v3", developerKey=YOUTUBE_API_KEY)
    query = " ".join(keywords)
    request = youtube.search().list(
        q=query,
        part="snippet",
        maxResults=25,
        type="video"
    )
    response = request.execute()
    videos = []
    for item in response["items"]:
        video = {
            "title": item["snippet"]["title"],
            "link": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
        }
        videos.append(video)
    return videos
    
@app.route("/", methods=["GET", "POST"])
def index():
    result, videos = None, []
    if request.method == "POST":
        text = ""
        if "input_text" in request.form and request.form["input_text"]:
            text = request.form["input_text"]
        elif "file" in request.files and request.files["file"]:
            text = extract_text_from_file(request.files["file"])
        elif "image" in request.files and request.files["image"]:
            text = extract_text_from_image(request.files["image"])

        if text:
            label, confidence = classify_text(text)
            keywords = extract_keywords(text)
            videos = search_youtube(keywords)
            result = {
                "text": text,
                "label": label,
                "confidence": confidence,
                "keywords": keywords
            }

    return render_template("index.html", result=result, videos=videos)
if __name__ == "__main__":
    app.run(debug=True)

"""